# app/services/file_service.py
import os
import uuid
import re
from pathlib import Path
from typing import Optional
import fitz  # PyMuPDF
from fastapi import UploadFile
from docx import Document

UPLOAD_DIR = "app/static/uploads"
ALLOWED_IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp"}
ALLOWED_PDF_EXTENSIONS = {".pdf"}
ALLOWED_RTF_EXTENSIONS = {".rtf"}
ALLOWED_DOC_EXTENSIONS = {".doc", ".docx"}

def ensure_upload_dir():
    """确保上传目录存在"""
    Path(UPLOAD_DIR).mkdir(parents=True, exist_ok=True)

async def process_uploaded_file(file: UploadFile, user_id: int) -> Optional[str]:
    """
    处理上传的文件
    - 如果是图片，直接保存
    - 如果是PDF，转换第一页为图片
    返回相对于static目录的文件路径
    """
    ensure_upload_dir()
    
    # 获取文件扩展名
    file_ext = Path(file.filename).suffix.lower()
    
    # 生成唯一文件名
    unique_filename = f"{uuid.uuid4()}{file_ext}"
    user_dir = Path(UPLOAD_DIR) / f"user_{user_id}"
    user_dir.mkdir(exist_ok=True)
    
    file_path = user_dir / unique_filename
    
    # 保存原始文件
    with open(file_path, "wb") as f:
        content = await file.read()
        f.write(content)
    
    # 如果是PDF，转换为图片
    if file_ext in ALLOWED_PDF_EXTENSIONS:
        try:
            # 使用PyMuPDF转换PDF第一页为图片
            doc = fitz.open(file_path)
            page = doc[0]  # 获取第一页
            mat = fitz.Matrix(2.0, 2.0)  # 缩放矩阵，提高分辨率
            pix = page.get_pixmap(matrix=mat)
            
            # 生成图片文件名
            img_filename = f"{uuid.uuid4()}.jpg"
            img_path = user_dir / img_filename
            
            # 保存为JPEG
            pix.save(str(img_path))
            doc.close()
            
            # 删除原始PDF文件
            os.remove(file_path)
            
            # 返回相对路径
            return f"/static/uploads/user_{user_id}/{img_filename}"
            
        except Exception as e:
            print(f"Error converting PDF to image: {e}")
            # 如果转换失败，删除文件并返回None
            if os.path.exists(file_path):
                os.remove(file_path)
            return None
    
    elif file_ext in ALLOWED_IMAGE_EXTENSIONS:
        # 直接返回图片路径
        return f"/static/uploads/user_{user_id}/{unique_filename}"
    
    elif file_ext in ALLOWED_RTF_EXTENSIONS:
        # 转换RTF为DOC并返回路径
        return convert_rtf_to_doc(file_path, user_id)
    
    else:
        # 不支持的文件类型，删除文件
        os.remove(file_path)
        return None

def get_file_info(file_path: str) -> dict:
    """获取文件信息"""
    if os.path.exists(file_path):
        stat = os.stat(file_path)
        return {
            "size": stat.st_size,
            "modified": stat.st_mtime,
            "exists": True
        }
    return {"exists": False}

def simple_rtf_to_text(rtf_content: str) -> str:
    """
    简单的RTF到纯文本转换功能
    移除RTF格式标记，提取纯文本内容
    """
    # 移除RTF头部信息
    text = re.sub(r'\\rtf\d+', '', rtf_content)
    
    # 移除字体表
    text = re.sub(r'\\fonttbl[^}]*}', '', text)
    
    # 移除颜色表
    text = re.sub(r'\\colortbl[^}]*}', '', text)
    
    # 移除样式表
    text = re.sub(r'\\stylesheet[^}]*}', '', text)
    
    # 移除其他RTF控制词
    text = re.sub(r'\\[a-z]+\d*\s?', '', text)
    
    # 移除RTF特殊字符
    text = re.sub(r'\\[{}\\]', '', text)
    
    # 移除多余的大括号
    text = re.sub(r'[{}]', '', text)
    
    # 清理多余的空白字符
    text = re.sub(r'\s+', ' ', text)
    
    return text.strip()

def convert_rtf_to_doc(rtf_path: str, user_id: int) -> Optional[str]:
    """
    将RTF文件转换为DOCX文件
    """
    try:
        # 读取RTF文件内容
        with open(rtf_path, 'r', encoding='utf-8', errors='ignore') as f:
            rtf_content = f.read()
        
        # 转换为纯文本
        text_content = simple_rtf_to_text(rtf_content)
        
        # 创建新的DOCX文档
        doc = Document()
        
        # 添加文档标题
        doc.add_heading('从RTF转换的文档', 0)
        
        # 将文本按段落分割并添加到文档中
        paragraphs = text_content.split('\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())
        
        # 生成新的文件名
        doc_filename = f"{uuid.uuid4()}.docx"
        user_dir = Path(UPLOAD_DIR) / f"user_{user_id}"
        user_dir.mkdir(exist_ok=True)
        doc_path = user_dir / doc_filename
        
        # 保存DOCX文件
        doc.save(str(doc_path))
        
        # 删除原始RTF文件
        os.remove(rtf_path)
        
        # 返回相对路径
        return f"/static/uploads/user_{user_id}/{doc_filename}"
        
    except Exception as e:
        print(f"Error converting RTF to DOC: {e}")
        # 如果转换失败，删除文件并返回None
        if os.path.exists(rtf_path):
            os.remove(rtf_path)
        return None
